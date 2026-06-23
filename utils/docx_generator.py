from docx import Document
from docx.shared import Pt, Inches
from io import BytesIO
import logging
from datetime import datetime
import re

class DocumentGenerator:
    @staticmethod
    def create_cover_letter_docx(cover_letter_text: str) -> BytesIO:
        """
        Generates an enterprise-formatted Word Document (.docx).
        Strictly enforces paragraph breaks regardless of how the LLM outputs \n.
        """
        logging.info("Generating professional DOCX file for Cover Letter...")
        
        doc = Document()
        
        # 1. Setup Official Business Margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # 2. Setup Default Formal Font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)

        # 3. Add Current Date
        current_date = datetime.now().strftime("%B %d, %Y")
        date_p = doc.add_paragraph(current_date)
        date_p.paragraph_format.space_after = Pt(24)

        # 4. Clean and Split Text into true paragraphs
        # Replace literal '\n' string with actual newline character, then split
        clean_text = cover_letter_text.replace('\\n', '\n')
        
        # Split by any form of line break
        paragraphs = re.split(r'\n+', clean_text)
        
        for para in paragraphs:
            para = para.strip()
            if para: # Only add if the paragraph is not empty
                p = doc.add_paragraph(para)
                p.paragraph_format.space_after = Pt(12) # Add space between paragraphs
                p.paragraph_format.line_spacing = 1.15

        # Save to buffer
        file_buffer = BytesIO()
        doc.save(file_buffer)
        file_buffer.seek(0)
        
        logging.info("DOCX generation complete with perfect paragraph formatting.")
        return file_buffer